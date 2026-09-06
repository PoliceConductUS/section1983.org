#!/usr/bin/env python3
"""
Build downloadable copies of the sample documents.

Reads the ```text blocks out of the sample articles and writes, for each:
  public/downloads/<slug>.txt   plain text, for pasting into any editor
  public/downloads/<slug>.docx  Word document, Courier New 12, one paragraph per line
  public/downloads/<slug>.pdf   rendered with headless Chrome from a monospace HTML page

Run:  python3 scripts/build-downloads.py
Requires python-docx and Google Chrome (macOS path below, or CHROME env var).
"""
import html
import os
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt, Inches
except ImportError:
    sys.exit("python-docx is required: pip install python-docx")

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "public" / "downloads"
CHROME = os.environ.get(
    "CHROME", "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"
)

# slug -> (article file, title, which text block(s) to include, 0-based)
SAMPLES = {
    "sample-section-1983-complaint-false-arrest": (
        "sample-section-1983-complaint-for-a-false-arrest-case.md",
        "Sample Section 1983 Complaint (False Arrest)",
        [0],
    ),
    "sample-declaration": (
        "sample-declaration-for-a-section-1983-case.md",
        "Sample Declaration (28 U.S.C. § 1746)",
        [0],
    ),
    "sample-response-to-motion-to-dismiss": (
        "sample-response-to-a-motion-to-dismiss-in-a-section-1983-case.md",
        "Sample Response to a Motion to Dismiss",
        [0],
    ),
    "sample-motion-for-extension-of-time": (
        "sample-motion-for-extension-of-time.md",
        "Sample Motion for Extension of Time, with Proposed Order",
        [0, 1],
    ),
    "sample-motion-for-leave-to-amend": (
        "sample-motion-for-leave-to-amend.md",
        "Sample Motion for Leave to Amend, with Proposed Order",
        [0, 1],
    ),
    "sample-meet-and-confer-letter": (
        "sample-meet-and-confer-letter-discovery.md",
        "Sample Meet-and-Confer Letter, with Conference Record",
        [0, 1],
    ),
}

NOTICE = (
    "section1983.org sample document. Educational template, not legal advice, "
    "not verified for your court. Replace every bracketed item. Check your "
    "district's local rules and your judge's standing orders before filing. "
    "Downloaded from https://www.section1983.org/downloads/"
)


def text_blocks(md: str) -> list[str]:
    return re.findall(r"^```text\n(.*?)^```", md, flags=re.S | re.M)


def write_txt(path: Path, title: str, body: str) -> None:
    path.write_text(f"{title}\n{'=' * len(title)}\n\n{NOTICE}\n\n{body.rstrip()}\n")


def write_docx(path: Path, title: str, body: str) -> None:
    doc = Document()
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(1)
        section.top_margin = section.bottom_margin = Inches(1)
    style = doc.styles["Normal"]
    style.font.name = "Courier New"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p = doc.add_paragraph()
    r = p.add_run(NOTICE)
    r.italic = True
    r.font.size = Pt(9)
    doc.add_paragraph("")
    for line in body.rstrip().split("\n"):
        doc.add_paragraph(line if line.strip() else "")
    doc.save(path)


def write_pdf(path: Path, title: str, body: str) -> None:
    if not Path(CHROME).exists():
        print(f"  skip pdf (no Chrome at {CHROME})")
        return
    page = f"""<!doctype html><html><head><meta charset="utf-8"><title>{html.escape(title)}</title>
<style>
@page {{ size: Letter; margin: 1in; }}
body {{ font-family: "Courier New", Courier, monospace; font-size: 11pt; color: #000; }}
.notice {{ font-size: 8.5pt; font-style: italic; color: #333; margin-bottom: 1.5em; }}
pre {{ white-space: pre-wrap; font: inherit; margin: 0; }}
</style></head><body>
<p class="notice">{html.escape(NOTICE)}</p>
<pre>{html.escape(body.rstrip())}</pre>
</body></html>"""
    with tempfile.TemporaryDirectory() as td:
        src = Path(td) / "doc.html"
        src.write_text(page)
        out = Path(td) / "doc.pdf"
        subprocess.run(
            [
                CHROME,
                "--headless=new",
                "--disable-gpu",
                "--no-pdf-header-footer",
                f"--print-to-pdf={out}",
                f"file://{src}",
            ],
            check=True,
            capture_output=True,
            timeout=120,
        )
        shutil.copyfile(out, path)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    for slug, (article, title, which) in SAMPLES.items():
        md = (ROOT / "articles" / article).read_text()
        blocks = text_blocks(md)
        chosen = [blocks[i] for i in which if i < len(blocks)]
        if not chosen:
            print(f"!! no text block found in {article}")
            continue
        body = "\n\n" .join(b.rstrip() for b in chosen)
        write_txt(OUT / f"{slug}.txt", title, body)
        write_docx(OUT / f"{slug}.docx", title, body)
        write_pdf(OUT / f"{slug}.pdf", title, body)
        print(f"built {slug}: {len(chosen)} block(s), {len(body.split())} words")


if __name__ == "__main__":
    main()
